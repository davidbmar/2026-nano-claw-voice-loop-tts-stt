"""Turning an uploaded file into something the assistant can be asked about.

Two jobs live here: pulling readable text out of a file, and handing that text
to the evidence platform so it becomes retrievable.

Extraction happens on this side rather than in the platform because the
platform speaks only ``POST /v1/ingest/text`` over HTTP — it has no multipart
route, and no PDF reader at all. Sending it plain text keeps the whole feature
inside one repository and one deployment.

The failure that matters most is the scanned PDF. A tax document photographed
or scanned has no text layer, and every extractor returns an empty string for
it. Ingesting that empty string would produce a document that exists, looks
fine in the list, and answers nothing — the worst possible outcome. So an empty
extraction is a loud, specific error and no row is ever written.
"""

from __future__ import annotations

import io
import logging
import os
import re
from typing import Any

import httpx

log = logging.getLogger("nano-claw.documents")

# 25 MB per file. The aiohttp application cap is set above this so the request
# body (file + multipart framing) fits, and this is the number the customer is
# actually held to.
MAX_DOCUMENT_BYTES = 25 * 1024 * 1024

# Below this many characters a "successful" extraction is not credible — a
# 40-page PDF that yields nine characters is a scan with a stray watermark, not
# a document. Small genuine text files are still fine: the check applies only
# to formats that carry a text layer separately from their rendering.
MIN_EXTRACTED_CHARS = 20

TEXT_SUFFIXES = (".txt", ".md", ".markdown", ".text")
PDF_SUFFIXES = (".pdf",)
DOCX_SUFFIXES = (".docx",)

ACCEPTED_DESCRIPTION = "PDF, Word (.docx), plain text or Markdown"

_WHITESPACE_RUN = re.compile(r"[ \t]+")
_BLANK_RUN = re.compile(r"\n{3,}")


class DocumentError(Exception):
    """An upload that cannot become a document, with the HTTP status to say so."""

    status = 422

    def __init__(self, message: str, *, status: int | None = None) -> None:
        super().__init__(message)
        self.message = message
        if status is not None:
            self.status = status


class UnsupportedDocument(DocumentError):
    status = 415


class DocumentTooLarge(DocumentError):
    status = 413


class NoTextLayer(DocumentError):
    """A file we can parse, that turns out to carry no readable text."""


def classify(filename: str) -> str:
    """Return ``text``/``pdf``/``docx`` for a filename, or refuse it.

    Dispatch is on the extension rather than the browser-supplied media type:
    the media type is attacker-controlled and routinely wrong (Windows sends
    ``application/octet-stream`` for .docx), while the extension is what the
    customer actually sees.
    """

    suffix = os.path.splitext(filename or "")[1].lower()
    if suffix in TEXT_SUFFIXES:
        return "text"
    if suffix in PDF_SUFFIXES:
        return "pdf"
    if suffix in DOCX_SUFFIXES:
        return "docx"
    raise UnsupportedDocument(
        f"We can't read {suffix or 'that file type'} yet. "
        f"Supported: {ACCEPTED_DESCRIPTION}."
    )


def tidy(text: str) -> str:
    """Normalise whitespace without destroying paragraph structure.

    Segmentation on the platform side is a 180-word window that never crosses a
    section boundary, so blank lines carry meaning and are preserved; runs of
    spaces and page-break padding do not and are collapsed.
    """

    unified = text.replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    collapsed = _WHITESPACE_RUN.sub(" ", unified)
    lines = [line.rstrip() for line in collapsed.split("\n")]
    return _BLANK_RUN.sub("\n\n", "\n".join(lines)).strip()


def _extract_text(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - deployment misconfiguration
        raise DocumentError(
            "PDF support is not installed on this server.", status=503
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # An empty-password decrypt covers the common "protected but not
            # really" case; anything else is a genuine refusal.
            try:
                reader.decrypt("")
            except Exception:  # noqa: BLE001 - pypdf raises several types here
                raise NoTextLayer(
                    "This PDF is password-protected, so we can't read it."
                ) from None
        pages = [(page.extract_text() or "") for page in reader.pages]
    except NoTextLayer:
        raise
    except Exception as exc:  # noqa: BLE001 - malformed files raise broadly
        raise DocumentError(f"This PDF could not be read ({exc}).") from exc
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document as open_docx
    except ImportError as exc:  # pragma: no cover - deployment misconfiguration
        raise DocumentError(
            "Word support is not installed on this server.", status=503
        ) from exc

    try:
        document = open_docx(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - malformed files raise broadly
        raise DocumentError(f"This Word file could not be read ({exc}).") from exc

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


_EXTRACTORS = {"text": _extract_text, "pdf": _extract_pdf, "docx": _extract_docx}


def extract(filename: str, data: bytes) -> tuple[str, str]:
    """Return ``(kind, text)`` for an uploaded file, or raise a DocumentError."""

    if len(data) > MAX_DOCUMENT_BYTES:
        raise DocumentTooLarge(
            f"That file is larger than the {MAX_DOCUMENT_BYTES // (1024 * 1024)} MB limit."
        )
    if not data:
        raise NoTextLayer("That file is empty.")

    kind = classify(filename)
    text = tidy(_EXTRACTORS[kind](data))

    if len(text) < MIN_EXTRACTED_CHARS:
        if kind == "pdf":
            raise NoTextLayer(
                "This PDF has no text layer — it looks like a scan or a photo. "
                "We can't read those yet; try exporting a text-based PDF."
            )
        raise NoTextLayer("We couldn't find any readable text in that file.")
    return kind, text


def derive_title(filename: str, data: bytes) -> str:
    """A human-readable title for the document list.

    A drag-dropped IRS form arrives named f8655.pdf, which tells the reader
    nothing. The PDF's own metadata usually says "Form 8655 (Rev. ...)", so
    prefer that when it looks like a real title, and fall back to the filename
    stem otherwise. Best effort by design: a title must never be the reason an
    upload fails.
    """

    stem = os.path.splitext(os.path.basename(filename or ""))[0][:200] or "document"
    if classify_or_none(filename) != "pdf":
        return stem
    try:
        from pypdf import PdfReader

        meta = PdfReader(io.BytesIO(data)).metadata
        title = str(meta.title).strip() if meta and meta.title else ""
    except Exception:  # noqa: BLE001 - metadata is a nicety, never a gate
        return stem
    # Reject junk metadata: empty, generic, suspiciously short, or the
    # authoring tool's filename echoed back.
    if (
        not title
        or len(title) < 6
        or title.lower() in {"untitled", "document", "pdf"}
        or title.lower() == stem.lower()
        or title.lower().endswith(".indd")
    ):
        return stem
    return title[:200]


def classify_or_none(filename: str) -> str | None:
    """`classify`, but None instead of an exception, for callers that probe."""

    try:
        return classify(filename)
    except UnsupportedDocument:
        return None


# ---- platform ingestion --------------------------------------------------


class PlatformClient:
    """The narrow slice of the evidence platform this feature needs.

    This is the first call to the platform from the Python side — retrieval
    still belongs entirely to the Node service. Ingestion is here because the
    file is here, and shipping bytes to another process only to have it POST
    them back would buy nothing.
    """

    def __init__(
        self,
        *,
        base_url: str | None = None,
        tenant_id: str | None = None,
        principal_id: str | None = None,
        timeout: float = 60.0,
    ) -> None:
        self.base_url = (
            base_url
            or os.environ.get("NANO_CLAW_INTELLIGENCE_URL")
            or "http://127.0.0.1:8000"
        ).rstrip("/")
        self.tenant_id = (
            tenant_id or os.environ.get("NANO_CLAW_INTELLIGENCE_TENANT") or "personal"
        )
        self.principal_id = (
            principal_id
            or os.environ.get("NANO_CLAW_INTELLIGENCE_PRINCIPAL")
            or "nano-claw"
        )
        # Ingestion is synchronous on the platform: parse, segment and index
        # all finish before it answers. A big document legitimately takes
        # tens of seconds, so this timeout is far longer than a retrieval's.
        self.timeout = timeout

    @staticmethod
    def ingest_timeout(characters: int, floor: float = 120.0) -> float:
        """How long to wait for a synchronous ingest of this much text.

        A fixed timeout cannot work here. The IRS 1040 instructions are 630k
        characters and took just over a minute to segment and index; a one-page
        W-2 takes under a second. A flat 60s silently turned the first genuinely
        large upload into a failure — and, because the platform finished the
        work anyway, into an indexed document the registry had no handle to.

        Measured at roughly 10k characters per second, then given a wide margin
        because the platform shares its database with live retrieval.
        """

        return min(900.0, max(floor, characters / 5_000))

    def _policy(self, permission: str) -> dict[str, Any]:
        return {
            "tenant_id": self.tenant_id,
            "principal_id": self.principal_id,
            "permissions": [permission],
        }

    async def ingest_text(
        self,
        *,
        document_key: str,
        title: str,
        text: str,
        collection_id: str,
        source_ref: str,
        metadata: dict[str, Any] | None = None,
    ) -> str:
        """Index a document and return the platform's document id.

        The request shape is exact: the platform's contracts are declared with
        ``extra="forbid"``, so one stray field is a 422 rather than a warning.
        """

        payload = {
            "tenant_id": self.tenant_id,
            "policy": self._policy("knowledge:ingest"),
            "source_ref": source_ref,
            "document_key": document_key,
            "title": title,
            "text": text,
            "collection_ids": [collection_id],
            "source_kind": "upload",
            "metadata": metadata or {},
        }
        timeout = self.ingest_timeout(len(text), floor=self.timeout)
        try:
            response = await self._post_ingest(payload, timeout)
        except httpx.TimeoutException:
            # The platform may well have finished after we stopped listening,
            # which is how the first large upload left an indexed document the
            # registry could not name. `document_key` is deterministic, so
            # re-sending the identical body returns the same document_id
            # without creating a second copy — the retry heals that case
            # instead of compounding it.
            log.warning("ingest timed out after %.0fs, retrying once", timeout)
            try:
                response = await self._post_ingest(payload, timeout * 2)
            except httpx.TimeoutException as exc:
                raise DocumentError(
                    "This document is taking longer to index than we allow. "
                    "It may still finish in the background — reload in a "
                    "minute, and re-add it if it has not appeared.",
                    status=504,
                ) from exc
        if response.status_code >= 400:
            raise DocumentError(
                f"The document service rejected this file "
                f"({response.status_code}: {response.text[:200]})",
                status=502,
            )
        document_id = response.json().get("document_id")
        if not document_id:
            raise DocumentError(
                "The document service accepted the file but returned no id.",
                status=502,
            )
        return str(document_id)

    async def _post_ingest(self, payload: dict[str, Any], timeout: float):
        async with httpx.AsyncClient(timeout=timeout) as client:
            return await client.post(f"{self.base_url}/v1/ingest/text", json=payload)

    async def delete_document(self, platform_document_id: str) -> bool:
        """Remove a document and everything derived from it.

        Returns ``True`` when the platform no longer holds it — including when
        it never did (404), so a purge that has already half-run can be retried
        without getting stuck.
        """

        async with httpx.AsyncClient(timeout=self.timeout) as client:
            response = await client.delete(
                f"{self.base_url}/v1/documents/{platform_document_id}",
                headers={
                    "X-Tenant-Id": self.tenant_id,
                    "X-Principal-Id": self.principal_id,
                    "X-Permissions": "knowledge:ingest",
                },
            )
        if response.status_code == 404:
            return True
        if response.status_code >= 400:
            log.warning(
                "platform delete failed for %s: %s %s",
                platform_document_id,
                response.status_code,
                response.text[:200],
            )
            return False
        return True
